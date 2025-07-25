import sqlite3
import pandas as pd
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches
import openpyxl
from tkinter import messagebox
from tkinter import filedialog as fd
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from PyPDF2 import PdfReader, PdfWriter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import io

# Conecta ao banco (ou cria, se não existir)
conn = sqlite3.connect("contadores.db")
cursor = conn.cursor()

# Garante que a tabela contadores exista com os novos campos
cursor.execute("""
CREATE TABLE IF NOT EXISTS contadores (
    cnpj          TEXT PRIMARY KEY,
    razao_social  TEXT NOT NULL,
    nome_fantasia TEXT NOT NULL,
    municipio     TEXT NOT NULL,
    contador      TEXT NOT NULL,
    tel_contador  TEXT NOT NULL,
    email         TEXT,
    endereco      TEXT,
    possui_certificado TEXT,
    situacao      TEXT
)
""")
conn.commit()

# Cria a tabela repis se não existir com os novos campos
conn_repis = sqlite3.connect("repis.db")
cursor_repis = conn_repis.cursor()

for col, col_def in [
    ("razao_social", "TEXT NOT NULL"),
    ("nome_fantasia", "TEXT"),
    # … liste aqui outras colunas que entraram depois
]:
    try:
        cursor_repis.execute(f"ALTER TABLE repis ADD COLUMN {col} {col_def}")
    except sqlite3.OperationalError:
        pass  # coluna já existe
conn_repis.commit()

cursor_repis.execute("""
CREATE TABLE IF NOT EXISTS repis (
    cnpj                TEXT PRIMARY KEY,
    razao_social        TEXT NOT NULL,
    nome_fantasia       TEXT,
    endereco            TEXT,
    complemento         TEXT,
    cep                 TEXT,
    email               TEXT,
    bairro              TEXT,
    uf                  TEXT,
    municipio           TEXT,
    data_abertura       TEXT,
    nome_solicitante    TEXT,
    solicitante_tipo    TEXT,
    telefone            TEXT,
    email_solicitante   TEXT,
    cpf                 TEXT,
    rg                  TEXT,
    contador            TEXT,
    telefone_contador   TEXT,
    email_contador      TEXT
)
""")
conn_repis.commit()

# Cria a tabela contadores_novo se não existir com os novos campos
conn_contadores_novo = sqlite3.connect("contadores_novo.db")
cursor_contadores_novo = conn_contadores_novo.cursor()

cursor_contadores_novo.execute("""
CREATE TABLE IF NOT EXISTS contadores_novo (
    cnpj                TEXT PRIMARY KEY,
    nome                TEXT NOT NULL,
    municipio           TEXT NOT NULL,
    socio               TEXT NOT NULL,
    contato             TEXT NOT NULL,
    tipo_pessoa         TEXT NOT NULL,
    empresas_representadas TEXT,
    empresa_associada_1 TEXT,
    empresa_associada_2 TEXT,
    empresa_associada_3 TEXT,
    email               TEXT
)""")
conn_contadores_novo.commit()

# Cria a tabela empresas se não existir
conn_empresas = sqlite3.connect("empresas.db")
cursor_empresas = conn_empresas.cursor()

cursor_empresas.execute("""
CREATE TABLE IF NOT EXISTS empresas (
    cnpj                TEXT PRIMARY KEY,
    razao_social        TEXT NOT NULL,
    nome_fantasia       TEXT,
    endereco            TEXT,
    complemento         TEXT,
    cep                 TEXT,
    email               TEXT,
    bairro              TEXT,
    uf                  TEXT,
    municipio           TEXT,
    telefone            TEXT,
    atividade_principal TEXT,
    data_abertura       TEXT,
    situacao            TEXT,
    responsavel         TEXT,
    telefone_responsavel TEXT,
    email_responsavel   TEXT
)
""")
conn_empresas.commit()

# Cria a tabela de associação contador-empresa
conn_relacionamentos = sqlite3.connect("relacionamentos.db")
cursor_relacionamentos = conn_relacionamentos.cursor()

cursor_relacionamentos.execute("""
CREATE TABLE IF NOT EXISTS contador_empresa (
    cnpj_contador TEXT NOT NULL,
    cnpj_empresa TEXT NOT NULL,
    tipo_relacao TEXT,
    ativo INTEGER DEFAULT 1,
    PRIMARY KEY (cnpj_contador, cnpj_empresa),
    FOREIGN KEY (cnpj_contador) REFERENCES contadores_novo(cnpj),
    FOREIGN KEY (cnpj_empresa) REFERENCES empresas(cnpj)
)
""")
conn_relacionamentos.commit()

# Cria a tabela de associação empresa-repis
cursor_relacionamentos.execute("""
CREATE TABLE IF NOT EXISTS empresa_repis (
    cnpj_empresa TEXT NOT NULL,
    cnpj_repis TEXT NOT NULL,
    ano_repis TEXT,
    cnpj_contador_responsavel TEXT,
    status TEXT,
    PRIMARY KEY (cnpj_empresa, cnpj_repis),
    FOREIGN KEY (cnpj_empresa) REFERENCES empresas(cnpj),
    FOREIGN KEY (cnpj_repis) REFERENCES repis(cnpj),
    FOREIGN KEY (cnpj_contador_responsavel) REFERENCES contadores_novo(cnpj)
)
""")
conn_relacionamentos.commit()


def criar_contador(dados):
    """
    Insere um novo registro de contador.
    Aceita uma lista com os dados: [cnpj, razao_social, nome_fantasia, municipio, contador, tel_contador, email, endereco, possui_certificado, situacao]
    """
    cnpj, razao_social, nome_fantasia, municipio, contador, tel_contador, email, endereco, possui_certificado, situacao = dados
    cursor.execute("""
        INSERT INTO contadores
          (cnpj, razao_social, nome_fantasia, municipio, contador, tel_contador, email, endereco, possui_certificado, situacao)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (cnpj, razao_social, nome_fantasia, municipio, contador, tel_contador, email, endereco, possui_certificado, situacao)
    )
    conn.commit()

def excluir_contador(cnpj):
    """
    Exclui o registro de contador com o CNPJ informado.
    """
    cursor.execute(
        "DELETE FROM contadores WHERE cnpj = ?",
        (cnpj,)
    )
    conn.commit()

def atualizar_contador(cnpj, *, razao_social=None, nome_fantasia=None,
                      municipio=None, contador=None, tel_contador=None, email=None, endereco=None, possui_certificado=None, situacao=None):
    """
    Atualiza apenas os campos não-None para o contador identificado pelo CNPJ.
    Exemplo:
      atualizar_contador("00.000.000/0001-00", razao_social="Nova Razão")
    """
    campos = []
    valores = []

    if razao_social is not None:
        campos.append("razao_social = ?")
        valores.append(razao_social)
    if nome_fantasia is not None:
        campos.append("nome_fantasia = ?")
        valores.append(nome_fantasia)
    if municipio is not None:
        campos.append("municipio = ?")
        valores.append(municipio)
    if contador is not None:
        campos.append("contador = ?")
        valores.append(contador)
    if tel_contador is not None:
        campos.append("tel_contador = ?")
        valores.append(tel_contador)
    if email is not None:
        campos.append("email = ?")
        valores.append(email)
    if endereco is not None:
        campos.append("endereco = ?")
        valores.append(endereco)
    if possui_certificado is not None:
        campos.append("possui_certificado = ?")
        valores.append(possui_certificado)
    if situacao is not None:
        campos.append("situacao = ?")
        valores.append(situacao)

    if not campos:
        # nada a atualizar
        return

    valores.append(cnpj)  # condição WHERE
    sql = f"UPDATE contadores SET {', '.join(campos)} WHERE cnpj = ?"
    cursor.execute(sql, valores)
    conn.commit()


def buscar_todos_contadores():
    """
    Retorna todos os registros de contadores do banco de dados.
    """
    cursor.execute("SELECT * FROM contadores")
    return cursor.fetchall()

def ver_dados():
    lista = []
    with conn:
        cur =conn.cursor()
        cur.execute('SELECT cnpj, razao_social, nome_fantasia, municipio, contador, tel_contador, email, endereco, possui_certificado, situacao FROM contadores')
        linha = cur.fetchall()

        for i in linha:
            lista.append(i)
    return lista

# Funções para a tabela repis com novos campos
def criar_repis(dados):
    cnpj, razao_social, nome_fantasia, endereco, complemento, cep, email, bairro, uf, municipio, data_abertura, nome_solicitante, solicitante_tipo, telefone, email_solicitante, cpf, rg, contador, telefone_contador, email_contador = dados
    cursor_repis.execute("""
        INSERT INTO repis
          (cnpj, razao_social, nome_fantasia, endereco, complemento, cep, email, bairro, uf, municipio, data_abertura, nome_solicitante, solicitante_tipo, telefone, email_solicitante, cpf, rg, contador, telefone_contador, email_contador)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (cnpj, razao_social, nome_fantasia, endereco, complemento, cep, email, bairro, uf, municipio, data_abertura, nome_solicitante, solicitante_tipo, telefone, email_solicitante, cpf, rg, contador, telefone_contador, email_contador)
    )
    conn_repis.commit()

def buscar_repis_por_cnpj(cnpj):
    cursor_repis.execute("SELECT * FROM repis WHERE cnpj = ?", (cnpj,))
    return cursor_repis.fetchone()

def atualizar_repis(cnpj, **kwargs):
    campos = []
    valores = []

    for campo, valor in kwargs.items():
        if valor is not None:
            campos.append(f"{campo} = ?")
            valores.append(valor)

    if not campos:
        return

    valores.append(cnpj)
    sql = f"UPDATE repis SET {', '.join(campos)} WHERE cnpj = ?"
    cursor_repis.execute(sql, valores)
    conn_repis.commit()

def excluir_repis(cnpj):
    cursor_repis.execute("DELETE FROM repis WHERE cnpj = ?", (cnpj,))
    conn_repis.commit()

def ver_dados_repis():
    lista = []
    with conn_repis:
        cur = conn_repis.cursor()
        cur.execute('SELECT cnpj, razao_social, nome_fantasia, endereco, complemento, cep, email, bairro, uf, municipio, data_abertura, nome_solicitante, solicitante_tipo, telefone, email_solicitante, cpf, rg, contador, telefone_contador, email_contador FROM repis')
        linha = cur.fetchall()

        for i in linha:
            lista.append(i)
    return lista

# Funções para a nova tabela contadores com novos campos
def criar_contador_novo(dados):
    cnpj, nome, municipio, socio, contato, tipo_pessoa, empresas_representadas, empresa_associada_1, empresa_associada_2, empresa_associada_3, email = dados
    cursor_contadores_novo.execute("""
        INSERT INTO contadores_novo
          (cnpj, nome, municipio, socio, contato, tipo_pessoa, empresas_representadas, empresa_associada_1, empresa_associada_2, empresa_associada_3, email)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (cnpj, nome, municipio, socio, contato, tipo_pessoa, empresas_representadas, empresa_associada_1, empresa_associada_2, empresa_associada_3, email)
    )
    conn_contadores_novo.commit()

def buscar_contador_por_cnpj_ou_nome(termo):
    cursor_contadores_novo.execute("SELECT * FROM contadores_novo WHERE cnpj = ? OR nome LIKE ?", (termo, f'%{termo}%'))
    return cursor_contadores_novo.fetchone()

def atualizar_contador_novo(cnpj, **kwargs):
    campos = []
    valores = []

    if 'email' in kwargs and kwargs['email'] is not None:
        campos.append("email = ?")
        valores.append(kwargs['email'])

    for campo, valor in kwargs.items():
        if campo != 'email' and valor is not None:
            campos.append(f"{campo} = ?")
            valores.append(valor)

    if not campos:
        return

    valores.append(cnpj)
    sql = f"UPDATE contadores_novo SET {", ".join(campos)} WHERE cnpj = ?"
    cursor_contadores_novo.execute(sql, valores)
    conn_contadores_novo.commit()

def excluir_contador_novo(cnpj):
    cursor_contadores_novo.execute("DELETE FROM contadores_novo WHERE cnpj = ?", (cnpj,))
    conn_contadores_novo.commit()

def ver_dados_contadores():
    lista = []
    with conn_contadores_novo:
        cur = conn_contadores_novo.cursor()
        cur.execute('SELECT cnpj, nome, municipio, socio, contato, email, tipo_pessoa, empresas_representadas, empresa_associada_1, empresa_associada_2, empresa_associada_3, email FROM contadores_novo')
        linha = cur.fetchall()

        for i in linha:
            lista.append(i)
    return lista

# Funções para a tabela empresas
def criar_empresa(dados):
    cnpj, razao_social, nome_fantasia, endereco, complemento, cep, email, bairro, uf, municipio, telefone, atividade_principal, data_abertura, situacao, responsavel, telefone_responsavel, email_responsavel = dados
    cursor_empresas.execute("""
        INSERT INTO empresas
          (cnpj, razao_social, nome_fantasia, endereco, complemento, cep, email, bairro, uf, municipio, telefone, atividade_principal, data_abertura, situacao, responsavel, telefone_responsavel, email_responsavel)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (cnpj, razao_social, nome_fantasia, endereco, complemento, cep, email, bairro, uf, municipio, telefone, atividade_principal, data_abertura, situacao, responsavel, telefone_responsavel, email_responsavel)
    )
    conn_empresas.commit()

def buscar_empresa_por_cnpj(cnpj):
    cursor_empresas.execute("SELECT * FROM empresas WHERE cnpj = ?", (cnpj,))
    return cursor_empresas.fetchone()

def atualizar_empresa(cnpj, **kwargs):
    campos = []
    valores = []

    for campo, valor in kwargs.items():
        if valor is not None:
            campos.append(f"{campo} = ?")
            valores.append(valor)

    if not campos:
        return

    valores.append(cnpj)
    sql = f"UPDATE empresas SET {', '.join(campos)} WHERE cnpj = ?"
    cursor_empresas.execute(sql, valores)
    conn_empresas.commit()

def excluir_empresa(cnpj):
    cursor_empresas.execute("DELETE FROM empresas WHERE cnpj = ?", (cnpj,))
    conn_empresas.commit()

def ver_dados_empresas():
    lista = []
    with conn_empresas:
        cur = conn_empresas.cursor()
        cur.execute('SELECT cnpj, razao_social, nome_fantasia, endereco, complemento, cep, email, bairro, uf, municipio, telefone, atividade_principal, data_abertura, situacao, responsavel, telefone_responsavel, email_responsavel FROM empresas')
        linha = cur.fetchall()

        for i in linha:
            lista.append(i)
    return lista

# Funções de exportação
def exportar_para_pdf(tabela):
    try:
        if tabela == "REPIS":
            dados = ver_dados_repis()
            headers = ["CNPJ", "Razão Social", "Nome Fantasia", "Endereço", "Complemento", "CEP", "E-mail", "Bairro", "UF", "Município", "Data Abertura", "Nome Solicitante", "Tipo Solicitante", "Telefone", "Email Solicitante", "CPF", "RG", "Contador", "Tel. Contador", "Email Contador"]
            nome_arquivo = "repis_dados.pdf"
        elif tabela == "Contadores":
            dados = ver_dados_contadores()
            headers = ["CNPJ", "Nome", "Município", "Sócio", "Contato", "email", "Tipo Pessoa", "Empresas Representadas", "Empresa Associada 1", "Empresa Associada 2", "Empresa Associada 3"]
            nome_arquivo = "contadores_dados.pdf"
        elif tabela == "Empresas":
            dados = ver_dados_empresas()
            headers = ["CNPJ", "Razão Social", "Nome Fantasia", "Endereço", "Complemento", "CEP", "E-mail", "Bairro", "UF", "Município", "Telefone", "Atividade Principal", "Data Abertura", "Situação", "Responsável", "Tel. Responsável", "Email Responsável"]
            nome_arquivo = "empresas_dados.pdf"
        
        doc = SimpleDocTemplate(nome_arquivo, pagesize=letter)
        elements = []
        
        # Título
        styles = getSampleStyleSheet()
        title = Paragraph(f"Relatório - {tabela}", styles["Title"])
        elements.append(title)
        
        # Tabela
        data = [headers] + dados
        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),     # Linha do cabeçalho
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 8),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),   # Corpo da tabela
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        elements.append(table)
        doc.build(elements)
        
        messagebox.showinfo('Sucesso', f'Arquivo {nome_arquivo} criado com sucesso!')
        
    except Exception as e:
        messagebox.showerror('Erro', f'Erro ao criar PDF: {str(e)}')

def exportar_para_word(tabela):
    try:
        if tabela == "REPIS":
            dados = ver_dados_repis()
            headers = ['CNPJ', 'Razão Social', 'Nome Fantasia', 'Endereço', 'Complemento', 'CEP', 'E-mail', 'Bairro', 'UF', 'Município', 'Data Abertura', 'Nome Solicitante', 'Tipo Solicitante', 'Telefone', 'Email Solicitante', 'CPF', 'RG', 'Contador', 'Tel. Contador', 'Email Contador']
            nome_arquivo = "repis_dados.docx"
        elif tabela == "Contadores":
            dados = ver_dados_contadores()
            headers = ["CNPJ", "Nome", "Município", "Sócio", "Contato", "email", "Tipo Pessoa", "Empresas Representadas", "Empresa Associada 1", "Empresa Associada 2", "Empresa Associada 3"]
            nome_arquivo = "contadores_dados.docx"
        else:  # Empresas
            dados = ver_dados_empresas()
            headers = ['CNPJ', 'Razão Social', 'Nome Fantasia', 'Endereço', 'Complemento', 'CEP', 'E-mail', 'Bairro', 'UF', 'Município', 'Telefone', 'Atividade Principal', 'Data Abertura', 'Situação', 'Responsável', 'Tel. Responsável', 'Email Responsável']
            nome_arquivo = "empresas_dados.docx"
        
        doc = Document()
        doc.add_heading(f'Relatório - {tabela}', 0)
        
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Cabeçalhos
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
        
        # Dados
        for row_data in dados:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data) if cell_data else ""
        
        doc.save(nome_arquivo)
        messagebox.showinfo('Sucesso', f'Arquivo {nome_arquivo} criado com sucesso!')
        
    except Exception as e:
        messagebox.showerror('Erro', f'Erro ao criar Word: {str(e)}')

def exportar_para_excel(tabela):
    try:
        if tabela == "REPIS":
            dados = ver_dados_repis()
            headers = ['CNPJ', 'Razão Social', 'Nome Fantasia', 'Endereço', 'Complemento', 'CEP', 'E-mail', 'Bairro', 'UF', 'Município', 'Data Abertura', 'Nome Solicitante', 'Tipo Solicitante', 'Telefone', 'Email Solicitante', 'CPF', 'RG', 'Contador', 'Tel. Contador', 'Email Contador']
            nome_arquivo = "repis_dados.xlsx"
        elif tabela == "Contadores":
            dados = ver_dados_contadores()
            headers = ["CNPJ", "Nome", "Município", "Sócio", "Contato", "email", "Tipo Pessoa", "Empresas Representadas", "Empresa Associada 1", "Empresa Associada 2", "Empresa Associada 3"]
            nome_arquivo = "contadores_dados.xlsx"
        else:  # Empresas
            dados = ver_dados_empresas()
            headers = ['CNPJ', 'Razão Social', 'Nome Fantasia', 'Endereço', 'Complemento', 'CEP', 'E-mail', 'Bairro', 'UF', 'Município', 'Telefone', 'Atividade Principal', 'Data Abertura', 'Situação', 'Responsável', 'Tel. Responsável', 'Email Responsável']
            nome_arquivo = "empresas_dados.xlsx"
        
        df = pd.DataFrame(dados, columns=headers)
        df.to_excel(nome_arquivo, index=False)
        
        messagebox.showinfo('Sucesso', f'Arquivo {nome_arquivo} criado com sucesso!')
        
    except Exception as e:
        messagebox.showerror('Erro', f'Erro ao criar Excel: {str(e)}')

# Cria a tabela de associação contador-empresa se não exções para gerenciar associação contador-empresa



# === FUNÇÕES DE INTERLIGAÇÃO DE TABELAS ===

# Funções para interligação de tabelas

def associar_contador_empresa(cnpj_contador, cnpj_empresa, tipo_relacao='representacao'):
    """Associa um contador a uma empresa"""
    try:
        conn_rel = sqlite3.connect('relacionamentos.db')
        cursor_rel = conn_rel.cursor()
        
        cursor_rel.execute("""
        INSERT OR REPLACE INTO contador_empresa 
        (cnpj_contador, cnpj_empresa, tipo_relacao, ativo)
        VALUES (?, ?, ?, 1)
        """, (cnpj_contador, cnpj_empresa, tipo_relacao))
        
        conn_rel.commit()
        conn_rel.close()
        return True
    except Exception as e:
        print(f"Erro ao associar contador-empresa: {e}")
        return False

def desassociar_contador_empresa(cnpj_contador, cnpj_empresa):
    """Desassocia um contador de uma empresa"""
    try:
        conn_rel = sqlite3.connect('relacionamentos.db')
        cursor_rel = conn_rel.cursor()
        
        cursor_rel.execute("""
        UPDATE contador_empresa 
        SET ativo = 0 
        WHERE cnpj_contador = ? AND cnpj_empresa = ?
        """, (cnpj_contador, cnpj_empresa))
        
        conn_rel.commit()
        conn_rel.close()
        return True
    except Exception as e:
        print(f"Erro ao desassociar contador-empresa: {e}")
        return False

def buscar_empresas_por_contador(cnpj_contador):
    """Busca todas as empresas associadas a um contador"""
    try:
        conn_rel = sqlite3.connect('relacionamentos.db')
        conn_empresas = sqlite3.connect('empresas.db')
        
        cursor_rel = conn_rel.cursor()
        cursor_empresas = conn_empresas.cursor()
        
        # Buscar CNPJs das empresas associadas
        cursor_rel.execute("""
        SELECT cnpj_empresa FROM contador_empresa 
        WHERE cnpj_contador = ? AND ativo = 1
        """, (cnpj_contador,))
        
        cnpjs_empresas = [row[0] for row in cursor_rel.fetchall()]
        
        empresas = []
        for cnpj in cnpjs_empresas:
            cursor_empresas.execute("SELECT * FROM empresas WHERE cnpj = ?", (cnpj,))
            empresa = cursor_empresas.fetchone()
            if empresa:
                empresas.append(empresa)
        
        conn_rel.close()
        conn_empresas.close()
        
        return empresas
    except Exception as e:
        print(f"Erro ao buscar empresas por contador: {e}")
        return []

def buscar_contadores_por_empresa(cnpj_empresa):
    """Busca todos os contadores associados a uma empresa"""
    try:
        conn_rel = sqlite3.connect('relacionamentos.db')
        conn_contadores = sqlite3.connect('contadores_novo.db')
        
        cursor_rel = conn_rel.cursor()
        cursor_contadores = conn_contadores.cursor()
        
        # Buscar CNPJs dos contadores associados
        cursor_rel.execute("""
        SELECT cnpj_contador FROM contador_empresa 
        WHERE cnpj_empresa = ? AND ativo = 1
        """, (cnpj_empresa,))
        
        cnpjs_contadores = [row[0] for row in cursor_rel.fetchall()]
        
        contadores = []
        for cnpj in cnpjs_contadores:
            cursor_contadores.execute("SELECT * FROM contadores_novo WHERE cnpj = ?", (cnpj,))
            contador = cursor_contadores.fetchone()
            if contador:
                contadores.append(contador)
        
        conn_rel.close()
        conn_contadores.close()
        
        return contadores
    except Exception as e:
        print(f"Erro ao buscar contadores por empresa: {e}")
        return []

def associar_empresa_repis(cnpj_empresa, cnpj_repis, cnpj_contador_responsavel=None, ano_repis='2025-2026'):
    """Associa uma empresa a um REPIS"""
    try:
        conn_rel = sqlite3.connect('relacionamentos.db')
        cursor_rel = conn_rel.cursor()
        
        cursor_rel.execute("""
        INSERT OR REPLACE INTO empresa_repis 
        (cnpj_empresa, cnpj_repis, ano_repis, cnpj_contador_responsavel, status)
        VALUES (?, ?, ?, ?, 'ativo')
        """, (cnpj_empresa, cnpj_repis, ano_repis, cnpj_contador_responsavel))
        
        conn_rel.commit()
        conn_rel.close()
        return True
    except Exception as e:
        print(f"Erro ao associar empresa-repis: {e}")
        return False

def buscar_repis_por_contador(cnpj_contador):
    """Busca todos os REPIS associados a um contador"""
    try:
        conn_rel = sqlite3.connect("relacionamentos.db")
        conn_repis = sqlite3.connect("repis.db")
        
        cursor_rel = conn_rel.cursor()
        cursor_repis = conn_repis.cursor()
        
        # Buscar CNPJs das empresas associadas ao contador através da tabela contador_empresa
        cursor_rel.execute("""
        SELECT cnpj_empresa FROM contador_empresa 
        WHERE cnpj_contador = ? AND ativo = 1
        """, (cnpj_contador,))
        
        cnpjs_empresas = [row[0] for row in cursor_rel.fetchall()]
        
        repis_list = []
        for cnpj_empresa in cnpjs_empresas:
            # Agora, buscar os REPIS associados a cada uma dessas empresas
            cursor_rel.execute("""
            SELECT cnpj_repis FROM empresa_repis
            WHERE cnpj_empresa = ? AND status = 'ativo'
            """, (cnpj_empresa,))
            
            cnpjs_repis_empresa = [row[0] for row in cursor_rel.fetchall()]
            
            for cnpj_repis in cnpjs_repis_empresa:
                cursor_repis.execute("SELECT * FROM repis WHERE cnpj = ?", (cnpj_repis,))
                repis = cursor_repis.fetchone()
                if repis:
                    repis_list.append(repis)
        
        conn_rel.close()
        conn_repis.close()
        
        return repis_list
    except Exception as e:
        print(f"Erro ao buscar REPIS por contador: {e}")
        return []

def buscar_dados_completos_contador(cnpj_contador):
    """Busca dados completos de um contador incluindo empresas e REPIS associados"""
    try:
        from view import buscar_contador_por_cnpj_ou_nome
        
        # Dados básicos do contador
        contador = buscar_contador_por_cnpj_ou_nome(cnpj_contador)
        if not contador:
            return None
        
        # Empresas associadas
        empresas = buscar_empresas_por_contador(cnpj_contador)
        
        # REPIS associados
        repis_list = buscar_repis_por_contador(cnpj_contador)
        
        return {
            'contador': contador,
            'empresas': empresas,
            'repis': repis_list
        }
    except Exception as e:
        print(f"Erro ao buscar dados completos: {e}")
        return None
